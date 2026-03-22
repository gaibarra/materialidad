from __future__ import annotations

import io
import json
import os
import re
import hashlib
from pathlib import Path
from decimal import Decimal
from types import SimpleNamespace
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.files.storage import default_storage
from django.utils import timezone
from django.utils.text import slugify

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm, inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
    Image,
)
from reportlab.graphics.shapes import Drawing, Rect, String, Circle
from reportlab.graphics.charts.piecharts import Pie

from .models import AuditMaterialityDossier, CompliancePillar, LegalConsultation, Operacion, OperacionEntregable
from .services import _detect_legal_consultation_focus, get_legal_consultation_type_label

__all__ = [
    "markdown_to_docx_bytes",
    "build_operacion_dossier_zip",
    "build_operacion_defensa_pdf",
    "build_audit_materiality_markdown",
    "build_audit_materiality_docx",
    "build_audit_materiality_pdf",
    "build_legal_consultation_pdf",
]

_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_PATTERN = re.compile(r"^\s*[-*+]\s+(.*)$")
_ORDERED_PATTERN = re.compile(r"^\s*(\d+)[\.)]\s+(.*)$")
_INLINE_TOKEN_PATTERN = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)")


def _append_formatted_runs(paragraph, text: str) -> None:
    """Convierte énfasis básico de Markdown en runs con estilo."""

    if not text:
        paragraph.add_run("")
        return

    for token in _INLINE_TOKEN_PATTERN.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _add_paragraph(document: Document, text: str, *, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _append_formatted_runs(paragraph, text)
    return paragraph


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Convierte Markdown sencillo a un archivo DOCX en memoria."""

    if not markdown_text or not markdown_text.strip():
        raise ValueError("El contenido del contrato no puede estar vacío")

    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            _add_paragraph(document, "")
            continue

        if stripped in {"---", "***", "___"}:
            _add_paragraph(document, "")
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            level = min(len(hashes), 4)
            heading = document.add_heading("", level=level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _append_formatted_runs(heading, text.strip())
            continue

        bullet_match = _BULLET_PATTERN.match(stripped)
        if bullet_match:
            _add_paragraph(document, bullet_match.group(1).strip(), style="List Bullet")
            continue

        ordered_match = _ORDERED_PATTERN.match(stripped)
        if ordered_match:
            _add_paragraph(document, ordered_match.group(2).strip(), style="List Number")
            continue

        _add_paragraph(document, stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _ts(dt) -> str:
    if not dt:
        return ""
    if hasattr(dt, "isoformat"):
        return dt.isoformat()
    return str(dt)


def _safe_path(prefix: str, label: str, extension: str | None = None) -> str:
    safe_label = slugify(label) or "item"
    ext = extension or ""
    if ext and not ext.startswith("."):
        ext = f".{ext}"
    return f"{prefix}/{safe_label}{ext}".strip("/")


def _add_readme(zf: ZipFile, index: dict[str, Any]) -> None:
    lines = [
        "Dossier de defensa de materialidad",
        f"Operacion ID: {index.get('operacion_id')}",
        f"Empresa: {index.get('empresa')}",
        f"Proveedor: {index.get('proveedor') or '-'}",
        f"Contrato: {index.get('contrato') or '-'}",
        f"Generado: {index.get('generated_at')}",
        "",
        "Entradas ordenadas cronologicamente por pilar SAT:",
    ]
    for entry in index.get("entries", []):
        ts = entry.get("timestamp") or ""
        pillar = entry.get("pillar") or ""
        title = entry.get("title") or entry.get("kind")
        ref = entry.get("file_path") or entry.get("external_url") or "(sin archivo)"
        lines.append(f"- [{ts}] ({pillar}) {title} -> {ref}")
    zf.writestr("README.txt", "\n".join(lines))


def _add_json_index(zf: ZipFile, index: dict[str, Any]) -> None:
    payload = json.dumps(index, ensure_ascii=False, indent=2)
    zf.writestr("indice.json", payload)


def _sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def _stream_file_to_zip(zf: ZipFile, storage_path: str, zip_path: str) -> bool:
    if not storage_path:
        return False
    if not default_storage.exists(storage_path):
        return False
    with default_storage.open(storage_path, "rb") as fh:
        zf.writestr(zip_path, fh.read())
    return True


def _collect_entregables(operacion: Operacion) -> Iterable[dict[str, Any]]:
    entregables = getattr(operacion, "entregables_prefetched", None)
    items = entregables if entregables is not None else operacion.entregables.all()
    for ent in items:
        yield {
            "kind": "entregable",
            "id": ent.id,
            "title": ent.titulo,
            "pillar": ent.pillar,
            "timestamp": _ts(ent.created_at),
            "metadata": {
                "estado": ent.estado,
                "tipo_gasto": ent.tipo_gasto,
                "codigo": ent.codigo,
                "fecha_compromiso": _ts(ent.fecha_compromiso),
                "fecha_entregado": _ts(ent.fecha_entregado),
                "fecha_recepcion": _ts(ent.fecha_recepcion),
                "fecha_factura": _ts(ent.fecha_factura),
            },
            "external_url": ent.oc_archivo_url or "",
        }


def _collect_evidencias(operacion: Operacion) -> Iterable[dict[str, Any]]:
    evidencias = getattr(operacion, "evidencias_prefetched", None)
    items = evidencias if evidencias is not None else operacion.evidencias.all()
    for ev in items:
        extension = os.path.splitext(ev.archivo.name or "")[1]
        yield {
            "kind": "evidencia",
            "id": ev.id,
            "title": ev.descripcion,
            "pillar": CompliancePillar.ENTREGABLES,
            "timestamp": _ts(ev.created_at),
            "file_name": ev.archivo.name,
            "zip_path": _safe_path("evidencias", ev.descripcion or f"evidencia-{ev.id}", extension),
            "metadata": {"tipo": ev.tipo},
        }


def _collect_contrato_archivos(operacion: Operacion) -> Iterable[dict[str, Any]]:
    contrato = operacion.contrato
    if not contrato:
        return []
    entries: list[dict[str, Any]] = []
    if contrato.archivo_notariado:
        extension = os.path.splitext(contrato.archivo_notariado.name or "")[1]
        entries.append(
            {
                "kind": "contrato",
                "id": contrato.id,
                "title": "fedatario_testimonio",
                "pillar": CompliancePillar.RAZON_NEGOCIO,
                "timestamp": _ts(contrato.fecha_ratificacion or contrato.updated_at),
                "file_name": contrato.archivo_notariado.name,
                "zip_path": _safe_path("contrato", "testimonio", extension),
                "metadata": {
                    "modalidad": contrato.firma_modalidad,
                    "fedatario": contrato.fedatario_nombre,
                    "numero_instrumento": contrato.numero_instrumento,
                },
            }
        )
    return entries


def build_operacion_dossier_zip(operacion: Operacion) -> bytes:
    """Construye un ZIP con indice para defender materialidad por operacion."""

    buffer = io.BytesIO()
    now = timezone.now().isoformat()

    entries: list[dict[str, Any]] = []
    entries.extend(_collect_entregables(operacion))
    entries.extend(_collect_evidencias(operacion))
    entries.extend(_collect_contrato_archivos(operacion))

    entries_sorted = sorted(entries, key=lambda e: e.get("timestamp") or "")
    index = {
        "operacion_id": operacion.id,
        "empresa": getattr(operacion.empresa, "razon_social", ""),
        "proveedor": getattr(operacion.proveedor, "razon_social", ""),
        "contrato": operacion.contrato_id,
        "monto": str(operacion.monto),
        "moneda": operacion.moneda,
        "fecha_operacion": _ts(operacion.fecha_operacion),
        "generated_at": now,
        "entries": entries_sorted,
    }

    with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as zf:
        manifest_files: list[dict[str, Any]] = []
        for entry in entries_sorted:
            file_name = entry.get("file_name")
            zip_path = entry.get("zip_path")
            if file_name and zip_path:
                file_bytes: bytes | None = None
                try:
                    with default_storage.open(file_name, "rb") as fh:
                        file_bytes = fh.read()
                except Exception:
                    file_bytes = None

                ok = file_bytes is not None
                if ok:
                    zf.writestr(zip_path, file_bytes)
                    entry["file_path"] = zip_path
                    manifest_files.append(
                        {
                            "path": zip_path,
                            "sha256": _sha256_bytes(file_bytes),
                            "size_bytes": len(file_bytes),
                            "kind": entry.get("kind"),
                            "entry_id": entry.get("id"),
                        }
                    )
                else:
                    entry["file_path"] = None
                    entry["missing_file"] = True

        index["entries"] = entries_sorted
        index_payload = json.dumps(index, ensure_ascii=False, indent=2).encode("utf-8")
        readme_lines = [
            "Dossier de defensa de materialidad",
            f"Operacion ID: {index.get('operacion_id')}",
            f"Empresa: {index.get('empresa')}",
            f"Proveedor: {index.get('proveedor') or '-'}",
            f"Contrato: {index.get('contrato') or '-'}",
            f"Generado: {index.get('generated_at')}",
            "",
            "Entradas ordenadas cronologicamente por pilar SAT:",
        ]
        for entry in index.get("entries", []):
            ts = entry.get("timestamp") or ""
            pillar = entry.get("pillar") or ""
            title = entry.get("title") or entry.get("kind")
            ref = entry.get("file_path") or entry.get("external_url") or "(sin archivo)"
            readme_lines.append(f"- [{ts}] ({pillar}) {title} -> {ref}")
        readme_payload = "\n".join(readme_lines).encode("utf-8")

        zf.writestr("indice.json", index_payload)
        zf.writestr("README.txt", readme_payload)

        manifest_files.append(
            {
                "path": "indice.json",
                "sha256": _sha256_bytes(index_payload),
                "size_bytes": len(index_payload),
                "kind": "meta",
                "entry_id": None,
            }
        )
        manifest_files.append(
            {
                "path": "README.txt",
                "sha256": _sha256_bytes(readme_payload),
                "size_bytes": len(readme_payload),
                "kind": "meta",
                "entry_id": None,
            }
        )

        manifest = {
            "operacion_id": operacion.id,
            "generated_at": now,
            "algorithm": "sha256",
            "files": manifest_files,
        }
        zf.writestr("manifiesto_integridad.json", json.dumps(manifest, ensure_ascii=False, indent=2))

    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Colores corporativos para el PDF profesional
# ---------------------------------------------------------------------------
_COLOR_PRIMARY = colors.HexColor("#1B2A4A")       # Azul marino oscuro
_COLOR_SECONDARY = colors.HexColor("#2E5090")      # Azul medio
_COLOR_ACCENT = colors.HexColor("#3B82F6")         # Azul brillante
_COLOR_LIGHT_BG = colors.HexColor("#F0F4FA")       # Fondo gris-azulado claro
_COLOR_BORDER = colors.HexColor("#CBD5E1")         # Gris borde
_COLOR_TEXT = colors.HexColor("#1E293B")           # Texto oscuro
_COLOR_TEXT_LIGHT = colors.HexColor("#64748B")     # Texto secundario
_COLOR_SUCCESS = colors.HexColor("#16A34A")        # Verde
_COLOR_WARNING = colors.HexColor("#F59E0B")        # Amarillo
_COLOR_DANGER = colors.HexColor("#DC2626")         # Rojo
_COLOR_WHITE = colors.white
_COLOR_TABLE_HEADER = colors.HexColor("#1E3A5F")   # Header de tablas


def _get_pdf_styles() -> dict[str, ParagraphStyle]:
    """Retorna estilos de párrafo profesionales para el reporte."""
    base = getSampleStyleSheet()
    return {
        "cover_title": ParagraphStyle(
            "CoverTitle", parent=base["Title"],
            fontName="Helvetica-Bold", fontSize=26, leading=32,
            textColor=_COLOR_WHITE, alignment=TA_CENTER, spaceAfter=6,
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitle", parent=base["Normal"],
            fontName="Helvetica", fontSize=13, leading=18,
            textColor=colors.HexColor("#CBD5E1"), alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "cover_detail": ParagraphStyle(
            "CoverDetail", parent=base["Normal"],
            fontName="Helvetica", fontSize=11, leading=15,
            textColor=colors.HexColor("#94A3B8"), alignment=TA_CENTER,
        ),
        "section_title": ParagraphStyle(
            "SectionTitle", parent=base["Heading1"],
            fontName="Helvetica-Bold", fontSize=16, leading=22,
            textColor=_COLOR_PRIMARY, spaceBefore=18, spaceAfter=10,
            borderPadding=(0, 0, 4, 0),
        ),
        "subsection_title": ParagraphStyle(
            "SubsectionTitle", parent=base["Heading2"],
            fontName="Helvetica-Bold", fontSize=12, leading=16,
            textColor=_COLOR_SECONDARY, spaceBefore=12, spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "BodyText_Custom", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=_COLOR_TEXT, alignment=TA_JUSTIFY, spaceAfter=6,
        ),
        "body_small": ParagraphStyle(
            "BodySmall", parent=base["Normal"],
            fontName="Helvetica", fontSize=8.5, leading=12,
            textColor=_COLOR_TEXT_LIGHT, spaceAfter=4,
        ),
        "table_header": ParagraphStyle(
            "TableHeader", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=12,
            textColor=_COLOR_WHITE, alignment=TA_LEFT,
        ),
        "table_cell": ParagraphStyle(
            "TableCell", parent=base["Normal"],
            fontName="Helvetica", fontSize=9, leading=12,
            textColor=_COLOR_TEXT, alignment=TA_LEFT,
        ),
        "table_cell_bold": ParagraphStyle(
            "TableCellBold", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=12,
            textColor=_COLOR_TEXT, alignment=TA_LEFT,
        ),
        "table_cell_right": ParagraphStyle(
            "TableCellRight", parent=base["Normal"],
            fontName="Helvetica", fontSize=9, leading=12,
            textColor=_COLOR_TEXT, alignment=TA_RIGHT,
        ),
        "footer": ParagraphStyle(
            "Footer", parent=base["Normal"],
            fontName="Helvetica", fontSize=7.5, leading=10,
            textColor=_COLOR_TEXT_LIGHT, alignment=TA_CENTER,
        ),
        "legal": ParagraphStyle(
            "Legal", parent=base["Normal"],
            fontName="Helvetica-Oblique", fontSize=8, leading=11,
            textColor=_COLOR_TEXT_LIGHT, alignment=TA_JUSTIFY,
            spaceBefore=4, spaceAfter=4,
        ),
        "badge_ok": ParagraphStyle(
            "BadgeOK", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=12,
            textColor=_COLOR_SUCCESS,
        ),
        "badge_warn": ParagraphStyle(
            "BadgeWarn", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=12,
            textColor=_COLOR_WARNING,
        ),
        "badge_danger": ParagraphStyle(
            "BadgeDanger", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=12,
            textColor=_COLOR_DANGER,
        ),
    }


def _format_currency(amount, currency: str = "MXN") -> str:
    """Formatea un monto como moneda."""
    try:
        val = Decimal(str(amount))
        formatted = f"${val:,.2f}"
    except Exception:
        formatted = str(amount)
    symbols = {"MXN": "MXN", "USD": "USD", "EUR": "EUR"}
    return f"{formatted} {symbols.get(currency, currency)}"


def _format_date_pdf(dt) -> str:
    """Formatea una fecha de forma legible."""
    if not dt:
        return "\u2014"
    if hasattr(dt, "strftime"):
        return dt.strftime("%d/%m/%Y")
    s = str(dt)
    if "T" in s:
        s = s.split("T")[0]
    parts = s.split("-")
    if len(parts) == 3:
        return f"{parts[2]}/{parts[1]}/{parts[0]}"
    return s


def _risk_color(nivel: str) -> colors.Color:
    """Color del semáforo de riesgo."""
    nivel_upper = (nivel or "").upper()
    if nivel_upper in ("BAJO", "LOW"):
        return _COLOR_SUCCESS
    if nivel_upper in ("MEDIO", "MEDIUM", "MODERADO"):
        return _COLOR_WARNING
    if nivel_upper in ("ALTO", "HIGH", "CRITICO"):
        return _COLOR_DANGER
    return _COLOR_TEXT_LIGHT


def _estatus_badge(estatus: str, styles: dict) -> Paragraph:
    """Retorna un Paragraph estilizado como badge de estatus."""
    estatus_upper = (estatus or "").upper()
    label_map = {
        "VALIDADO": ("Validado", styles["badge_ok"]),
        "EN_PROCESO": ("En proceso", styles["badge_warn"]),
        "PENDIENTE": ("Pendiente", styles["badge_warn"]),
        "RECHAZADO": ("Rechazado", styles["badge_danger"]),
        "ENTREGADO": ("Entregado", styles["badge_ok"]),
        "RECIBIDO": ("Recibido", styles["badge_ok"]),
        "FACTURADO": ("Facturado", styles["badge_ok"]),
        "VALIDO": ("Valido", styles["badge_ok"]),
        "INVALIDO": ("Invalido", styles["badge_danger"]),
        "NO_ENCONTRADO": ("No encontrado", styles["badge_danger"]),
        "VALIDADA": ("Validada", styles["badge_ok"]),
        "OBSERVADA": ("Observada", styles["badge_warn"]),
        "SIN_COINCIDENCIA": ("Sin coincidencia", styles["badge_ok"]),
        "PRESUNTO": ("Presunto", styles["badge_warn"]),
        "DEFINITIVO": ("Definitivo", styles["badge_danger"]),
        "APROBADO": ("Aprobado", styles["badge_ok"]),
        "EN_PROCESO": ("En proceso", styles["badge_warn"]),
    }
    label, style = label_map.get(estatus_upper, (estatus or "\u2014", styles["table_cell"]))
    return Paragraph(f"<b>{label}</b>", style)


def _build_risk_indicator(nivel: str, score: str) -> Drawing:
    """Dibuja un indicador visual de riesgo (semáforo circular)."""
    d = Drawing(180, 50)
    color = _risk_color(nivel)

    # Círculo principal
    circle = Circle(25, 25, 18)
    circle.fillColor = color
    circle.strokeColor = colors.HexColor("#E2E8F0")
    circle.strokeWidth = 2
    d.add(circle)

    # Texto del nivel dentro del círculo
    nivel_short = (nivel or "N/A")[:4].upper()
    label = String(25, 20, nivel_short, fontName="Helvetica-Bold",
                   fontSize=9, fillColor=_COLOR_WHITE, textAnchor="middle")
    d.add(label)

    # Score al lado
    score_text = f"Score: {score}" if score and score != "N/A" else ""
    if score_text:
        score_label = String(55, 28, score_text, fontName="Helvetica-Bold",
                             fontSize=10, fillColor=_COLOR_TEXT, textAnchor="start")
        d.add(score_label)

    risk_label = String(55, 14, f"Riesgo {nivel or 'N/A'}",
                        fontName="Helvetica", fontSize=9,
                        fillColor=_COLOR_TEXT_LIGHT, textAnchor="start")
    d.add(risk_label)

    return d


def _header_footer_factory(operacion: Operacion, empresa_name: str):
    """Retorna una función para dibujar header/footer en páginas de contenido."""

    def _draw_header_footer(canvas, doc):
        canvas.saveState()
        width, height = letter

        # --- Header ---
        canvas.setStrokeColor(_COLOR_BORDER)
        canvas.setLineWidth(0.5)
        canvas.line(1.5 * cm, height - 1.35 * cm, width - 1.5 * cm, height - 1.35 * cm)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(_COLOR_PRIMARY)
        canvas.drawString(1.5 * cm, height - 1.05 * cm,
                          f"REPORTE DE DEFENSA FISCAL  -  Operacion #{operacion.id}")

        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(_COLOR_TEXT_LIGHT)
        canvas.drawRightString(width - 1.5 * cm, height - 1.05 * cm,
                               empresa_name[:60])

        # --- Footer ---
        canvas.line(1.5 * cm, 1.8 * cm, width - 1.5 * cm, 1.8 * cm)

        canvas.setFont("Helvetica", 7)
        canvas.drawString(1.5 * cm, 1.2 * cm,
                          "CONFIDENCIAL - Documento generado para defensa de materialidad fiscal")
        canvas.drawRightString(width - 1.5 * cm, 1.2 * cm,
                               f"Pagina {doc.page}")
        canvas.setFont("Helvetica", 6.5)
        canvas.drawCentredString(width / 2, 0.7 * cm,
                                 f"Generado: {timezone.now().strftime('%d/%m/%Y %H:%M:%S')}  -  "
                                 "Art. 5-A CFF  -  Art. 69-B CFF  -  Materialidad")

        canvas.restoreState()

    return _draw_header_footer


def _build_cover_page(operacion: Operacion, styles: dict) -> list:
    """Construye la primera hoja ejecutiva del reporte."""
    elements: list = []
    width = letter[0] - 3 * cm

    riesgo = (operacion.metadata or {}).get("riesgo_materialidad") or {}
    riesgo_nivel = str(riesgo.get("nivel", "N/A"))
    riesgo_score = str(riesgo.get("score", "N/A"))
    empresa_name = getattr(operacion.empresa, "razon_social", "") or "\u2014"
    proveedor_name = getattr(operacion.proveedor, "razon_social", "") or "\u2014"
    empresa_rfc = getattr(operacion.empresa, "rfc", "") or "\u2014"
    proveedor_rfc = getattr(operacion.proveedor, "rfc", "") or "\u2014"
    contrato_nombre = operacion.contrato.nombre if operacion.contrato else "Sin contrato asignado"
    mark_path = _institutional_mark_path()

    elements.append(Spacer(1, 0.45 * cm))

    brand_cell: Any
    if mark_path:
        brand_cell = Table([
            [
                Image(mark_path, width=1.2 * cm, height=1.2 * cm),
                Paragraph(
                    "<b>Materialidad Legal Fiscal</b><br/><font size='8' color='#64748B'>Emision institucional</font>",
                    styles["table_cell"],
                ),
            ]
        ], colWidths=[1.6 * cm, 4.6 * cm])
        brand_cell.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
    else:
        brand_cell = Paragraph(
            "<b>Materialidad Legal Fiscal</b><br/><font size='8' color='#64748B'>Emision institucional</font>",
            styles["table_cell"],
        )

    executive_header = Table([
        [
            brand_cell,
            Paragraph(
                f"<font size='17'><b>REPORTE DE DEFENSA FISCAL</b></font><br/><font size='9' color='#64748B'>Sintesis ejecutiva de materialidad operativa. Operacion ID: {operacion.id}.</font>",
                styles["table_cell"],
            ),
        ]
    ], colWidths=[6.5 * cm, width - 6.5 * cm])
    executive_header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROUNDEDCORNERS", [8, 8, 8, 8]),
    ]))
    elements.append(executive_header)
    elements.append(Spacer(1, 10))

    summary_rows = [
        [
            Paragraph("<b>Operacion ID</b>", styles["table_cell_bold"]),
            Paragraph(str(operacion.id), styles["table_cell"]),
            Paragraph("<b>Fecha</b>", styles["table_cell_bold"]),
            Paragraph(_format_date_pdf(operacion.fecha_operacion), styles["table_cell"]),
        ],
        [
            Paragraph("<b>Monto</b>", styles["table_cell_bold"]),
            Paragraph(_format_currency(operacion.monto, operacion.moneda), styles["table_cell"]),
            Paragraph("<b>Estatus</b>", styles["table_cell_bold"]),
            _estatus_badge(operacion.estatus_validacion, styles),
        ],
        [
            Paragraph("<b>Riesgo</b>", styles["table_cell_bold"]),
            Paragraph(f"{riesgo_nivel} · Score {riesgo_score}", styles["table_cell"]),
            Paragraph("<b>Clasificación</b>", styles["table_cell_bold"]),
            Paragraph("Confidencial - Uso interno", styles["table_cell"]),
        ],
    ]
    summary_table = Table(summary_rows, colWidths=[2.5 * cm, 4.9 * cm, 2.8 * cm, width - 10.2 * cm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), _COLOR_LIGHT_BG),
        ("BACKGROUND", (2, 0), (2, -1), _COLOR_LIGHT_BG),
        ("ROWBACKGROUNDS", (1, 0), (1, -1), [_COLOR_WHITE, colors.HexColor("#F8FAFC")]),
        ("ROWBACKGROUNDS", (3, 0), (3, -1), [_COLOR_WHITE, colors.HexColor("#F8FAFC")]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 8))

    col_w = width / 2 - 4
    info_data = [
        [Paragraph("<b>EMPRESA</b>", styles["body_small"]),
         Paragraph("<b>PROVEEDOR</b>", styles["body_small"])],
        [Paragraph(empresa_name, styles["body"]),
         Paragraph(proveedor_name, styles["body"])],
        [Paragraph(f"RFC: {empresa_rfc}", styles["body_small"]),
         Paragraph(f"RFC: {proveedor_rfc}", styles["body_small"])],
    ]
    info_table = Table(info_data, colWidths=[col_w, col_w])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_LIGHT_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("LINEBELOW", (0, 0), (-1, 0), 1, _COLOR_BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 8))

    detail_data = [
        [Paragraph("<b>Contrato</b>", styles["body_small"]),
         Paragraph(contrato_nombre, styles["body"])],
        [Paragraph("<b>Tipo Operacion</b>", styles["body_small"]),
         Paragraph(operacion.tipo_operacion or "\u2014", styles["body"])],
        [Paragraph("<b>Concepto</b>", styles["body_small"]),
         Paragraph((operacion.concepto or "\u2014")[:200], styles["body"])],
        [Paragraph("<b>Moneda</b>", styles["body_small"]),
         Paragraph(operacion.moneda or "MXN", styles["body"])],
    ]
    detail_table = Table(detail_data, colWidths=[3.5 * cm, width - 3.5 * cm])
    detail_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), _COLOR_LIGHT_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, _COLOR_BORDER),
        ("BOX", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(detail_table)
    elements.append(Spacer(1, 8))

    info_note = Table([
        [
            Paragraph(
                "<b>Alcance</b><br/><font size='9'>Documento ejecutivo de defensa fiscal para revision directiva, fiscal y de auditoria. Debe leerse junto con el expediente digital y sus anexos.</font>",
                styles["table_cell"],
            ),
            Paragraph(
                f"<b>Generado</b><br/><font size='9'>{timezone.now().strftime('%d/%m/%Y %H:%M hrs.')}<br/>Art. 5-A CFF · Art. 69-B CFF</font>",
                styles["table_cell"],
            ),
        ]
    ], colWidths=[width * 0.62, width * 0.38])
    info_note.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(info_note)

    return elements


def _build_section_datos_operacion(operacion: Operacion, styles: dict) -> list:
    """Seccion 1: Datos generales de la operacion."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Paragraph("1. DATOS DE LA OPERACION", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    empresa = operacion.empresa
    proveedor = operacion.proveedor

    # Datos empresa
    elements.append(Paragraph("1.1 Empresa", styles["subsection_title"]))
    emp_rows = [
        [Paragraph("<b>Campo</b>", styles["table_header"]),
         Paragraph("<b>Valor</b>", styles["table_header"])],
        [Paragraph("Razon Social", styles["table_cell_bold"]),
         Paragraph(getattr(empresa, "razon_social", "") or "\u2014", styles["table_cell"])],
        [Paragraph("RFC", styles["table_cell_bold"]),
         Paragraph(getattr(empresa, "rfc", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Tipo Persona", styles["table_cell_bold"]),
         Paragraph(getattr(empresa, "tipo_persona", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Regimen Fiscal", styles["table_cell_bold"]),
         Paragraph(getattr(empresa, "regimen_fiscal", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Actividad Economica", styles["table_cell_bold"]),
         Paragraph(getattr(empresa, "actividad_economica", "") or "\u2014", styles["table_cell"])],
    ]
    domicilio = getattr(empresa, "domicilio_fiscal", None)
    if domicilio:
        emp_rows.append([
            Paragraph("Domicilio Fiscal", styles["table_cell_bold"]),
            Paragraph(str(domicilio), styles["table_cell"]),
        ])

    emp_table = Table(emp_rows, colWidths=[4.5 * cm, width - 4.5 * cm])
    emp_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("BACKGROUND", (0, 1), (-1, -1), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(emp_table)
    elements.append(Spacer(1, 12))

    # Datos proveedor
    elements.append(Paragraph("1.2 Proveedor", styles["subsection_title"]))
    prov_rows = [
        [Paragraph("<b>Campo</b>", styles["table_header"]),
         Paragraph("<b>Valor</b>", styles["table_header"])],
        [Paragraph("Razon Social", styles["table_cell_bold"]),
         Paragraph(getattr(proveedor, "razon_social", "") or "\u2014", styles["table_cell"])],
        [Paragraph("RFC", styles["table_cell_bold"]),
         Paragraph(getattr(proveedor, "rfc", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Tipo Persona", styles["table_cell_bold"]),
         Paragraph(getattr(proveedor, "tipo_persona", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Actividad Principal", styles["table_cell_bold"]),
         Paragraph(getattr(proveedor, "actividad_principal", "") or "\u2014", styles["table_cell"])],
        [Paragraph("Regimen Fiscal", styles["table_cell_bold"]),
         Paragraph(getattr(proveedor, "regimen_fiscal", "") or "\u2014", styles["table_cell"])],
    ]

    # Estatus SAT del proveedor
    estatus_sat = getattr(proveedor, "estatus_sat", "") or "\u2014"
    estatus_69b = getattr(proveedor, "estatus_69b", "") or "\u2014"
    riesgo_fiscal = getattr(proveedor, "riesgo_fiscal", "") or "\u2014"

    prov_rows.append([
        Paragraph("Estatus SAT", styles["table_cell_bold"]),
        Paragraph(estatus_sat, styles["table_cell"]),
    ])
    prov_rows.append([
        Paragraph("Estatus 69-B", styles["table_cell_bold"]),
        _estatus_badge(estatus_69b, styles) if estatus_69b != "\u2014" else
        Paragraph("\u2014", styles["table_cell"]),
    ])

    risk_style_key = (
        "badge_ok" if riesgo_fiscal.upper() == "BAJO" else
        "badge_warn" if riesgo_fiscal.upper() == "MEDIO" else
        "badge_danger" if riesgo_fiscal.upper() == "ALTO" else
        "table_cell"
    )
    prov_rows.append([
        Paragraph("Riesgo Fiscal", styles["table_cell_bold"]),
        Paragraph(f"<b>{riesgo_fiscal}</b>", styles[risk_style_key]),
    ])

    prov_table = Table(prov_rows, colWidths=[4.5 * cm, width - 4.5 * cm])
    prov_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("BACKGROUND", (0, 1), (-1, -1), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(prov_table)

    return elements


def _build_section_contrato(operacion: Operacion, styles: dict) -> list:
    """Seccion 2: Informacion del contrato."""
    elements: list = []
    width = letter[0] - 3 * cm
    contrato = operacion.contrato

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("2. CONTRATO VIGENTE", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    if not contrato:
        elements.append(Paragraph(
            "No se ha asociado un contrato a esta operacion. Se recomienda vincular un "
            "contrato para fortalecer la defensa de materialidad.",
            styles["body"],
        ))
        return elements

    rows = [
        [Paragraph("<b>Atributo</b>", styles["table_header"]),
         Paragraph("<b>Detalle</b>", styles["table_header"])],
        [Paragraph("Nombre", styles["table_cell_bold"]),
         Paragraph(contrato.nombre or "\u2014", styles["table_cell"])],
        [Paragraph("Codigo Interno", styles["table_cell_bold"]),
         Paragraph(contrato.codigo_interno or "\u2014", styles["table_cell"])],
        [Paragraph("Categoria", styles["table_cell_bold"]),
         Paragraph(contrato.categoria or "\u2014", styles["table_cell"])],
        [Paragraph("Proceso", styles["table_cell_bold"]),
         Paragraph(contrato.proceso or "\u2014", styles["table_cell"])],
        [Paragraph("Fecha de Firma", styles["table_cell_bold"]),
         Paragraph(_format_date_pdf(contrato.fecha_firma), styles["table_cell"])],
        [Paragraph("Vigencia", styles["table_cell_bold"]),
         Paragraph(
             f"{_format_date_pdf(contrato.vigencia_inicio)} - {_format_date_pdf(contrato.vigencia_fin)}",
             styles["table_cell"],
         )],
    ]

    # Razón de negocio
    if contrato.razon_negocio:
        rows.append([
            Paragraph("Razon de Negocio", styles["table_cell_bold"]),
            Paragraph(contrato.razon_negocio[:500], styles["table_cell"]),
        ])

    rows.append([
        Paragraph("Razon Negocio Estado", styles["table_cell_bold"]),
        _estatus_badge(contrato.razon_negocio_estado or "PENDIENTE", styles),
    ])

    # Fecha cierta
    fc_label = "Si" if contrato.fecha_cierta_obtenida else "No"
    rows.append([
        Paragraph("Fecha Cierta Obtenida", styles["table_cell_bold"]),
        Paragraph(f"<b>{fc_label}</b>", styles[
            "badge_ok" if contrato.fecha_cierta_obtenida else "badge_warn"
        ]),
    ])

    if contrato.firma_modalidad:
        rows.append([
            Paragraph("Modalidad de Firma", styles["table_cell_bold"]),
            Paragraph(contrato.firma_modalidad, styles["table_cell"]),
        ])

    if contrato.fedatario_nombre:
        rows.append([
            Paragraph("Fedatario", styles["table_cell_bold"]),
            Paragraph(contrato.fedatario_nombre, styles["table_cell"]),
        ])

    if contrato.numero_instrumento:
        rows.append([
            Paragraph("No. Instrumento", styles["table_cell_bold"]),
            Paragraph(contrato.numero_instrumento, styles["table_cell"]),
        ])

    if contrato.fecha_ratificacion:
        rows.append([
            Paragraph("Fecha Ratificacion", styles["table_cell_bold"]),
            Paragraph(_format_date_pdf(contrato.fecha_ratificacion), styles["table_cell"]),
        ])

    c_table = Table(rows, colWidths=[5 * cm, width - 5 * cm])
    c_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(c_table)

    return elements


def _build_section_entregables(operacion: Operacion, styles: dict) -> list:
    """Seccion 3: Entregables por pilar de cumplimiento."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("3. ENTREGABLES POR PILAR DE CUMPLIMIENTO", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    entregables_data = list(_collect_entregables(operacion))

    if not entregables_data:
        elements.append(Paragraph(
            "No se han registrado entregables para esta operacion.",
            styles["body"],
        ))
        return elements

    # Agrupar por pilar
    pilares: dict[str, list] = {}
    for ent in entregables_data:
        pilar = ent.get("pillar") or "SIN PILAR"
        pilares.setdefault(pilar, []).append(ent)

    pilar_labels = {
        CompliancePillar.ENTREGABLES: "Entregables",
        CompliancePillar.RAZON_NEGOCIO: "Razon de Negocio",
        CompliancePillar.CAPACIDAD_PROVEEDOR: "Capacidad del Proveedor",
        CompliancePillar.FECHA_CIERTA: "Fecha Cierta",
    }

    for pilar, items in pilares.items():
        pilar_name = pilar_labels.get(pilar, str(pilar))
        elements.append(Paragraph(pilar_name, styles["subsection_title"]))

        header = [
            Paragraph("<b>#</b>", styles["table_header"]),
            Paragraph("<b>Titulo</b>", styles["table_header"]),
            Paragraph("<b>Estado</b>", styles["table_header"]),
            Paragraph("<b>F. Compromiso</b>", styles["table_header"]),
            Paragraph("<b>F. Entregado</b>", styles["table_header"]),
        ]

        rows = [header]
        for i, ent in enumerate(items, 1):
            meta = ent.get("metadata", {})
            rows.append([
                Paragraph(str(i), styles["table_cell"]),
                Paragraph(ent.get("title", "\u2014")[:80], styles["table_cell"]),
                _estatus_badge(meta.get("estado", "PENDIENTE"), styles),
                Paragraph(_format_date_pdf(meta.get("fecha_compromiso")),
                          styles["table_cell"]),
                Paragraph(_format_date_pdf(meta.get("fecha_entregado")),
                          styles["table_cell"]),
            ])

        col_widths = [1 * cm, width - 9.5 * cm, 2.8 * cm, 2.8 * cm, 2.8 * cm]
        e_table = Table(rows, colWidths=col_widths, repeatRows=1)
        e_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
            ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
            ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
            ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(e_table)
        elements.append(Spacer(1, 8))

    return elements


def _build_section_evidencias(operacion: Operacion, styles: dict) -> list:
    """Seccion 4: Evidencias materiales."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("4. EVIDENCIAS MATERIALES", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    evidencias_data = list(_collect_evidencias(operacion))

    if not evidencias_data:
        elements.append(Paragraph(
            "No se han cargado evidencias materiales para esta operacion.",
            styles["body"],
        ))
        return elements

    header = [
        Paragraph("<b>#</b>", styles["table_header"]),
        Paragraph("<b>Descripcion</b>", styles["table_header"]),
        Paragraph("<b>Tipo</b>", styles["table_header"]),
        Paragraph("<b>Fecha</b>", styles["table_header"]),
    ]
    rows = [header]
    for i, ev in enumerate(evidencias_data, 1):
        rows.append([
            Paragraph(str(i), styles["table_cell"]),
            Paragraph(ev.get("title", "\u2014")[:120], styles["table_cell"]),
            Paragraph((ev.get("metadata", {}).get("tipo") or "\u2014"), styles["table_cell"]),
            Paragraph(ev.get("timestamp", "\u2014")[:10] if ev.get("timestamp") else "\u2014",
                      styles["table_cell"]),
        ])

    col_widths = [1 * cm, width - 6.5 * cm, 3 * cm, 2.5 * cm]
    ev_table = Table(rows, colWidths=col_widths, repeatRows=1)
    ev_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(ev_table)

    return elements


def _build_section_validaciones(operacion: Operacion, styles: dict) -> list:
    """Seccion 5: Estado de validaciones fiscales."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("5. VALIDACIONES FISCALES", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    riesgo = (operacion.metadata or {}).get("riesgo_materialidad") or {}
    riesgo_nivel = str(riesgo.get("nivel", "N/A"))
    riesgo_score = str(riesgo.get("score", "N/A"))

    # Indicador de riesgo visual
    elements.append(Paragraph("5.1 Evaluacion de Riesgo de Materialidad", styles["subsection_title"]))
    elements.append(_build_risk_indicator(riesgo_nivel, riesgo_score))
    elements.append(Spacer(1, 10))

    # Tabla de validaciones
    elements.append(Paragraph("5.2 Estado de Validaciones", styles["subsection_title"]))

    val_rows = [
        [Paragraph("<b>Validacion</b>", styles["table_header"]),
         Paragraph("<b>Estatus</b>", styles["table_header"]),
         Paragraph("<b>Ultima Verificacion</b>", styles["table_header"])],
        [Paragraph("Validacion General", styles["table_cell_bold"]),
         _estatus_badge(operacion.estatus_validacion, styles),
         Paragraph(_format_date_pdf(operacion.ultima_validacion), styles["table_cell"])],
        [Paragraph("CFDI", styles["table_cell_bold"]),
         _estatus_badge(operacion.cfdi_estatus or "PENDIENTE", styles),
         Paragraph(_format_date_pdf(operacion.ultima_validacion_cfdi), styles["table_cell"])],
        [Paragraph("SPEI", styles["table_cell_bold"]),
         _estatus_badge(operacion.spei_estatus or "PENDIENTE", styles),
         Paragraph(_format_date_pdf(operacion.ultima_validacion_spei), styles["table_cell"])],
    ]

    if operacion.uuid_cfdi:
        val_rows.append([
            Paragraph("UUID CFDI", styles["table_cell_bold"]),
            Paragraph(operacion.uuid_cfdi, styles["table_cell"]),
            Paragraph("\u2014", styles["table_cell"]),
        ])

    if operacion.referencia_spei:
        val_rows.append([
            Paragraph("Ref. SPEI", styles["table_cell_bold"]),
            Paragraph(operacion.referencia_spei, styles["table_cell"]),
            Paragraph("\u2014", styles["table_cell"]),
        ])

    if operacion.nif_aplicable:
        val_rows.append([
            Paragraph("NIF Aplicable", styles["table_cell_bold"]),
            Paragraph(operacion.nif_aplicable, styles["table_cell"]),
            Paragraph("\u2014", styles["table_cell"]),
        ])

    v_table = Table(val_rows, colWidths=[5 * cm, 5 * cm, width - 10 * cm], repeatRows=1)
    v_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(v_table)

    # Observaciones contables
    if operacion.observacion_contable:
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("5.3 Observacion Contable", styles["subsection_title"]))
        elements.append(Paragraph(operacion.observacion_contable[:1000], styles["body"]))

    return elements


def _build_section_indice_anexos(operacion: Operacion, styles: dict) -> list:
    """Seccion 6: Indice de anexos."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("6. INDICE DE ANEXOS", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph(
        "Los siguientes documentos conforman el expediente de defensa de materialidad "
        "y se encuentran disponibles en el dossier digital descargable (formato ZIP).",
        styles["body"],
    ))
    elements.append(Spacer(1, 6))

    anexos: list[tuple[str, str, str]] = []  # (tipo, titulo, pilar)

    for entry in _collect_entregables(operacion):
        anexos.append((
            "Entregable",
            entry.get("title") or "\u2014",
            str(entry.get("pillar") or "\u2014"),
        ))

    for entry in _collect_evidencias(operacion):
        anexos.append((
            "Evidencia",
            entry.get("title") or "\u2014",
            "Material",
        ))

    for entry in _collect_contrato_archivos(operacion):
        anexos.append((
            "Contrato",
            entry.get("title") or "\u2014",
            str(entry.get("pillar") or "\u2014"),
        ))

    if not anexos:
        elements.append(Paragraph(
            "No se han registrado anexos para esta operacion. "
            "Se recomienda cargar evidencia documental para fortalecer la defensa.",
            styles["body"],
        ))
        return elements

    header = [
        Paragraph("<b>#</b>", styles["table_header"]),
        Paragraph("<b>Tipo</b>", styles["table_header"]),
        Paragraph("<b>Descripcion</b>", styles["table_header"]),
        Paragraph("<b>Pilar</b>", styles["table_header"]),
    ]
    rows = [header]
    for i, (tipo, titulo, pilar) in enumerate(anexos, 1):
        rows.append([
            Paragraph(str(i), styles["table_cell"]),
            Paragraph(tipo, styles["table_cell_bold"]),
            Paragraph(titulo[:100], styles["table_cell"]),
            Paragraph(pilar, styles["table_cell"]),
        ])

    col_widths = [1 * cm, 2.5 * cm, width - 7 * cm, 3.5 * cm]
    a_table = Table(rows, colWidths=col_widths, repeatRows=1)
    a_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_TABLE_HEADER),
        ("TEXTCOLOR", (0, 0), (-1, 0), _COLOR_WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_COLOR_WHITE, _COLOR_LIGHT_BG]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(a_table)

    return elements


def _build_section_fundamento_legal(styles: dict) -> list:
    """Seccion 7: Fundamento legal."""
    elements: list = []

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("7. FUNDAMENTO LEGAL", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph(
        "El presente reporte se fundamenta en las siguientes disposiciones del "
        "Codigo Fiscal de la Federacion (CFF) y su normatividad secundaria:",
        styles["body"],
    ))
    elements.append(Spacer(1, 6))

    fundamentos = [
        (
            "Art. 5-A del CFF - Razon de Negocio",
            "Los actos juridicos que carezcan de una razon de negocio y que generen un beneficio "
            "fiscal directo o indirecto, tendran los efectos fiscales que correspondan a los que "
            "se habrian realizado para la obtencion del beneficio economico razonablemente esperado "
            "por el contribuyente."
        ),
        (
            "Art. 69-B del CFF - Operaciones Inexistentes",
            "La autoridad fiscal podra presumir la inexistencia de operaciones amparadas en "
            "comprobantes fiscales cuando detecte que un contribuyente ha estado emitiendo "
            "comprobantes sin contar con los activos, personal, infraestructura o capacidad "
            "material para prestar los servicios o producir, comercializar o entregar los bienes "
            "que amparan dichos comprobantes."
        ),
        (
            "Art. 69-B Bis del CFF - Transmision Indebida de Perdidas",
            "Cuando la autoridad fiscal detecte que una persona fisica o moral fue parte de una "
            "reestructuracion, escision, fusion, operacion o acto juridico, careciendo de una "
            "razon de negocio y obteniendo un beneficio fiscal, sera considerada como transmision "
            "indebida de perdidas fiscales."
        ),
        (
            "Resolucion Miscelanea Fiscal Vigente",
            "Las reglas de caracter general emitidas por el SAT que establecen los requisitos "
            "y procedimientos para acreditar la materialidad de las operaciones, incluyendo la "
            "documentacion soporte requerida."
        ),
    ]

    for titulo, texto in fundamentos:
        elements.append(Paragraph(f"<b>{titulo}</b>", styles["subsection_title"]))
        elements.append(Paragraph(texto, styles["legal"]))
        elements.append(Spacer(1, 4))

    return elements


def _build_closing(styles: dict) -> list:
    """Seccion de cierre del reporte."""
    elements: list = []
    width = letter[0] - 3 * cm

    elements.append(Spacer(1, 1 * cm))
    elements.append(HRFlowable(width="100%", thickness=2, color=_COLOR_PRIMARY))
    elements.append(Spacer(1, 10))

    # Caja de cierre
    close_data = [
        [Paragraph(
            "<b>FIN DEL REPORTE DE DEFENSA FISCAL</b>",
            ParagraphStyle("CloseTitle", fontName="Helvetica-Bold", fontSize=12,
                           textColor=_COLOR_PRIMARY, alignment=TA_CENTER),
        )],
        [Paragraph(
            "Este documento fue generado automaticamente por el sistema de Materialidad Fiscal. "
            "La informacion contenida proviene de los registros capturados y validados en la plataforma. "
            "Para cualquier aclaracion, consultar con el responsable del despacho.",
            ParagraphStyle("CloseBody", fontName="Helvetica", fontSize=9, leading=13,
                           textColor=_COLOR_TEXT_LIGHT, alignment=TA_CENTER),
        )],
        [Paragraph(
            f"Generado: {timezone.now().strftime('%d/%m/%Y %H:%M:%S')}",
            ParagraphStyle("CloseDate", fontName="Helvetica", fontSize=8,
                           textColor=colors.HexColor("#64748B"), alignment=TA_CENTER),
        )],
    ]

    close_table = Table(close_data, colWidths=[width])
    close_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (0, 0), 14),
        ("BOTTOMPADDING", (0, -1), (0, -1), 14),
        ("LEFTPADDING", (0, 0), (-1, -1), 20),
        ("RIGHTPADDING", (0, 0), (-1, -1), 20),
        ("TOPPADDING", (0, 1), (-1, 1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 4),
        ("ROUNDEDCORNERS", [8, 8, 8, 8]),
    ]))
    elements.append(close_table)

    return elements


def build_operacion_defensa_pdf(operacion: Operacion) -> bytes:
    """Genera un PDF profesional de defensa de materialidad fiscal."""

    buffer = io.BytesIO()
    styles = _get_pdf_styles()
    empresa_name = getattr(operacion.empresa, "razon_social", "") or "\u2014"
    page_width, page_height = letter

    hf_func = _header_footer_factory(operacion, empresa_name)

    content_frame = Frame(
        1.5 * cm, 2.05 * cm,
        page_width - 3 * cm, page_height - 4.0 * cm,
        id="content_frame",
    )

    doc = BaseDocTemplate(
        buffer,
        pagesize=letter,
        title=f"Reporte de Defensa Fiscal - Operacion #{operacion.id}",
        author="Sistema de Materialidad Fiscal",
        subject="Defensa de Materialidad",
    )

    doc.addPageTemplates([
        PageTemplate(id="content", frames=[content_frame], onPage=hf_func),
    ])

    # Construir elementos del documento
    elements: list = []

    elements.extend(_build_cover_page(operacion, styles))

    # Secciones
    elements.extend(_build_section_datos_operacion(operacion, styles))
    elements.extend(_build_section_contrato(operacion, styles))
    elements.extend(_build_section_entregables(operacion, styles))
    elements.extend(_build_section_evidencias(operacion, styles))
    elements.extend(_build_section_validaciones(operacion, styles))
    elements.extend(_build_section_indice_anexos(operacion, styles))
    elements.extend(_build_section_fundamento_legal(styles))
    elements.extend(_build_closing(styles))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def build_audit_materiality_markdown(dossier: AuditMaterialityDossier) -> str:
    payload = dossier.payload or {}
    benchmark_input = payload.get("benchmarkInput") or {}
    findings = payload.get("findings") or []
    empresa = getattr(dossier.empresa, "razon_social", "\u2014") or "\u2014"
    ejercicio = dossier.ejercicio
    mg = payload.get("mg")
    met = payload.get("met")
    trivial = payload.get("clearlyTrivial")
    benchmark_key = payload.get("selectedBenchmarkKey") or "sin-base"
    editor_name = dossier.last_edited_by_name or dossier.last_edited_by_email or "Sistema"
    edited_at = _format_date_pdf(dossier.updated_at)

    lines = [
        "# Expediente de Materialidad de Auditoría",
        "",
        f"**Empresa:** {empresa}",
        f"**Ejercicio:** {ejercicio}",
        f"**Última edición:** {editor_name} — {edited_at}",
        "",
        "## Criterio documentado",
        "",
        f"- **Base seleccionada:** {benchmark_key}",
        f"- **Materialidad global (MG):** {_format_currency(mg or 0)}",
        f"- **Materialidad de ejecución (MET):** {_format_currency(met or 0)}",
        f"- **Claramente trivial:** {_format_currency(trivial or 0)}",
        "",
        "## Bases capturadas",
        "",
        f"- Utilidad antes de impuestos: {_format_currency(benchmark_input.get('utilidadAntesImpuestos') or 0)}",
        f"- Ingresos: {_format_currency(benchmark_input.get('ingresos') or 0)}",
        f"- Activos totales: {_format_currency(benchmark_input.get('activos') or 0)}",
        f"- Capital contable: {_format_currency(benchmark_input.get('capital') or 0)}",
        f"- Gasto total: {_format_currency(benchmark_input.get('gastos') or 0)}",
        "",
        "## Hallazgos acumulados",
        "",
    ]

    if findings:
        for idx, finding in enumerate(findings, start=1):
            lines.extend(
                [
                    f"### {idx}. {finding.get('titulo') or 'Hallazgo sin título'}",
                    "",
                    f"- **Área:** {finding.get('area') or '\u2014'}",
                    f"- **Tipo de impacto:** {finding.get('impactoTipo') or '\u2014'}",
                    f"- **Severidad:** {finding.get('severidad') or '\u2014'}",
                    f"- **Impacto estimado:** {_format_currency(finding.get('impactoMonto') or 0)}",
                    f"- **Descripción:** {finding.get('descripcion') or '\u2014'}",
                    f"- **Próximo paso:** {finding.get('recomendacion') or '\u2014'}",
                    "",
                ]
            )
    else:
        lines.extend(["No hay hallazgos acumulados para este expediente.", ""])

    lines.extend(
        [
            "## Nota metodológica",
            "",
            "Este expediente resume la aplicación práctica de NIA 320 y NIA 450, priorizando la selección de una base defendible, la fijación de MG/MET y la acumulación de hallazgos cuantificables y cualitativos.",
            "",
        ]
    )
    return "\n".join(lines)


def build_audit_materiality_docx(dossier: AuditMaterialityDossier) -> bytes:
    return markdown_to_docx_bytes(build_audit_materiality_markdown(dossier))


def build_audit_materiality_pdf(dossier: AuditMaterialityDossier) -> bytes:
    payload = dossier.payload or {}
    benchmark_input = payload.get("benchmarkInput") or {}
    findings = payload.get("findings") or []
    styles = _get_pdf_styles()
    buffer = io.BytesIO()
    page_width, page_height = letter

    def _cover_page_template(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(_COLOR_PRIMARY)
        canvas.rect(0, page_height - 3, page_width, 3, fill=1, stroke=0)
        canvas.setFillColor(_COLOR_PRIMARY)
        canvas.rect(0, 0, page_width, 3, fill=1, stroke=0)
        canvas.restoreState()

    def _header_footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#E2E8F0"))
        canvas.setLineWidth(0.5)
        canvas.line(1.5 * cm, page_height - 1.5 * cm, page_width - 1.5 * cm, page_height - 1.5 * cm)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawString(1.5 * cm, 1.1 * cm, f"Materialidad de Auditoría · Empresa {getattr(dossier.empresa, 'rfc', '')}")
        canvas.drawRightString(page_width - 1.5 * cm, 1.1 * cm, f"Página {canvas.getPageNumber()}")
        canvas.restoreState()

    cover_frame = Frame(1.5 * cm, 1.5 * cm, page_width - 3 * cm, page_height - 3 * cm, id="cover_frame")
    content_frame = Frame(1.5 * cm, 2.2 * cm, page_width - 3 * cm, page_height - 4.2 * cm, id="content_frame")
    doc = BaseDocTemplate(
        buffer,
        pagesize=letter,
        title=f"Expediente de Materialidad de Auditoría - {dossier.empresa_id}/{dossier.ejercicio}",
        author="Sistema de Materialidad Fiscal",
        subject="Expediente de Materialidad de Auditoría",
    )
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[cover_frame], onPage=_cover_page_template),
        PageTemplate(id="content", frames=[content_frame], onPage=_header_footer),
    ])

    elements: list[Any] = []
    empresa = getattr(dossier.empresa, "razon_social", "\u2014") or "\u2014"
    editor_name = dossier.last_edited_by_name or dossier.last_edited_by_email or "Sistema"

    elements.append(Spacer(1, 2.3 * cm))
    elements.append(Paragraph("EXPEDIENTE DE MATERIALIDAD DE AUDITORÍA", styles["cover_title"]))
    elements.append(Spacer(1, 0.35 * cm))
    elements.append(Paragraph(f"{empresa} · Ejercicio {dossier.ejercicio}", styles["cover_subtitle"]))
    elements.append(Spacer(1, 0.5 * cm))
    elements.append(Paragraph(f"Última edición: {editor_name} · {_format_date_pdf(dossier.updated_at)}", styles["body"]))
    elements.append(NextPageTemplate("content"))
    elements.append(PageBreak())

    elements.append(Paragraph("Resumen ejecutivo", styles["section_title"]))
    elements.append(Paragraph(
        "Documento de trabajo para soportar la determinación de materialidad de auditoría, la base seleccionada y la acumulación de hallazgos conforme a NIA 320 y NIA 450.",
        styles["body"],
    ))
    elements.append(Spacer(1, 0.25 * cm))

    summary_rows = [
        [Paragraph("Empresa", styles["table_header"]), Paragraph(empresa, styles["table_cell"])],
        [Paragraph("Ejercicio", styles["table_header"]), Paragraph(str(dossier.ejercicio), styles["table_cell"])],
        [Paragraph("Base seleccionada", styles["table_header"]), Paragraph(str(payload.get("selectedBenchmarkKey") or "\u2014"), styles["table_cell"])],
        [Paragraph("MG", styles["table_header"]), Paragraph(_format_currency(payload.get("mg") or 0), styles["table_cell"])],
        [Paragraph("MET", styles["table_header"]), Paragraph(_format_currency(payload.get("met") or 0), styles["table_cell"])],
        [Paragraph("Claramente trivial", styles["table_header"]), Paragraph(_format_currency(payload.get("clearlyTrivial") or 0), styles["table_cell"])],
    ]
    summary_table = Table(summary_rows, colWidths=[5.2 * cm, 10.8 * cm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.white),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.35 * cm))

    elements.append(Paragraph("Bases capturadas", styles["section_title"]))
    base_rows = [
        [Paragraph("Utilidad antes de impuestos", styles["table_header"]), Paragraph(_format_currency(benchmark_input.get("utilidadAntesImpuestos") or 0), styles["table_cell"])],
        [Paragraph("Ingresos", styles["table_header"]), Paragraph(_format_currency(benchmark_input.get("ingresos") or 0), styles["table_cell"])],
        [Paragraph("Activos totales", styles["table_header"]), Paragraph(_format_currency(benchmark_input.get("activos") or 0), styles["table_cell"])],
        [Paragraph("Capital contable", styles["table_header"]), Paragraph(_format_currency(benchmark_input.get("capital") or 0), styles["table_cell"])],
        [Paragraph("Gasto total", styles["table_header"]), Paragraph(_format_currency(benchmark_input.get("gastos") or 0), styles["table_cell"])],
    ]
    base_table = Table(base_rows, colWidths=[6.1 * cm, 9.9 * cm])
    base_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(base_table)
    elements.append(Spacer(1, 0.35 * cm))

    elements.append(Paragraph("Hallazgos acumulados", styles["section_title"]))
    if findings:
        for idx, finding in enumerate(findings, start=1):
            impact_type = finding.get("impactoTipo") or "\u2014"
            severity = finding.get("severidad") or "\u2014"
            elements.append(Paragraph(f"{idx}. {finding.get('titulo') or 'Hallazgo sin título'}", styles["subsection_title"]))
            detail_rows = [
                [Paragraph("Área", styles["table_header"]), Paragraph(str(finding.get("area") or "\u2014"), styles["table_cell"])],
                [Paragraph("Tipo de impacto", styles["table_header"]), Paragraph(str(impact_type), styles["table_cell"])],
                [Paragraph("Severidad", styles["table_header"]), Paragraph(str(severity), styles["table_cell"])],
                [Paragraph("Impacto estimado", styles["table_header"]), Paragraph(_format_currency(finding.get("impactoMonto") or 0), styles["table_cell"])],
            ]
            detail_table = Table(detail_rows, colWidths=[4.8 * cm, 11.2 * cm])
            detail_table.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]))
            elements.append(detail_table)
            if finding.get("descripcion"):
                elements.append(Paragraph(f"<b>Descripción:</b> {finding.get('descripcion')}", styles["body"]))
            if finding.get("recomendacion"):
                elements.append(Paragraph(f"<b>Próximo paso:</b> {finding.get('recomendacion')}", styles["body"]))
            elements.append(Spacer(1, 0.25 * cm))
    else:
        elements.append(Paragraph("No hay hallazgos acumulados para este expediente.", styles["body"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def _legal_consultation_header_footer_factory(consultation_type_label: str):
    def _draw_header_footer(canvas, doc):
        canvas.saveState()
        width, height = letter
        canvas.setStrokeColor(_COLOR_BORDER)
        canvas.setLineWidth(0.5)
        canvas.line(1.5 * cm, height - 1.35 * cm, width - 1.5 * cm, height - 1.35 * cm)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(_COLOR_PRIMARY)
        canvas.drawString(1.5 * cm, height - 1.05 * cm, "OPINION LEGAL PRELIMINAR")
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(_COLOR_TEXT_LIGHT)
        canvas.drawRightString(width - 1.5 * cm, height - 1.05 * cm, consultation_type_label[:70])

        canvas.line(1.5 * cm, 1.6 * cm, width - 1.5 * cm, 1.6 * cm)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(1.5 * cm, 1.05 * cm, "CONFIDENCIAL - Materialidad Legal Fiscal - Uso interno del despacho y del Consejo")
        canvas.drawRightString(width - 1.5 * cm, 1.05 * cm, f"Pagina {doc.page}")
        canvas.restoreState()

    return _draw_header_footer


def _sanitize_pdf_text(value: str | None) -> str:
    normalized = re.sub(r"\s+", " ", (value or "")).strip()
    return escape(normalized) if normalized else "\u2014"


def _institutional_mark_path() -> str | None:
    candidate = Path(__file__).resolve().parents[2] / "frontend" / "public" / "icon-192x192.png"
    return str(candidate) if candidate.exists() else None


def _consultation_signatory_role(consultation: LegalConsultation) -> str:
    user = consultation.user
    if getattr(user, "is_superuser", False) or getattr(user, "is_staff", False):
        return "Socio responsable del despacho"
    return "Responsable de criterio fiscal"


def _markdown_to_pdf_flowables(markdown_text: str, styles: dict[str, ParagraphStyle]) -> list[Any]:
    flowables: list[Any] = []
    for raw_line in (markdown_text or "").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            flowables.append(Spacer(1, 4))
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            style_name = "section_title" if level <= 2 else "subsection_title"
            flowables.append(Paragraph(_sanitize_pdf_text(heading_match.group(2)), styles[style_name]))
            continue

        bullet_match = _BULLET_PATTERN.match(stripped)
        if bullet_match:
            flowables.append(Paragraph(f"• {_sanitize_pdf_text(bullet_match.group(1))}", styles["body"]))
            continue

        ordered_match = _ORDERED_PATTERN.match(stripped)
        if ordered_match:
            flowables.append(Paragraph(f"{ordered_match.group(1)}. {_sanitize_pdf_text(ordered_match.group(2))}", styles["body"]))
            continue

        plain_text = re.sub(r"^[>#`\s]+", "", stripped)
        flowables.append(Paragraph(_sanitize_pdf_text(plain_text), styles["body"]))

    return flowables


def build_legal_consultation_pdf(consultation: LegalConsultation) -> bytes:
    styles = _get_pdf_styles()
    buffer = io.BytesIO()
    references_payload = consultation.references if isinstance(consultation.references, list) else []
    reference_proxies = [
        SimpleNamespace(
            ley=ref.get("ley", ""),
            ordenamiento=ref.get("ordenamiento", ""),
            resumen=ref.get("resumen", ""),
            contenido=ref.get("extracto", ""),
            sat_categoria=ref.get("sat_categoria", ""),
            rubro=ref.get("rubro", ""),
            header=ref.get("header", ""),
        )
        for ref in references_payload
        if isinstance(ref, dict)
    ]
    focus = _detect_legal_consultation_focus(
        question=consultation.question,
        context_block=consultation.context or "",
        references=reference_proxies,
    )
    consultation_type_label = get_legal_consultation_type_label(focus)
    current_refs = [
        ref for ref in references_payload
        if isinstance(ref, dict) and (ref.get("es_vigente") or ref.get("estatus_vigencia") == "VIGENTE")
    ]
    support_label = (
        "Sustento vigente verificado" if current_refs else
        "Soporte histórico o por validar" if references_payload else
        "Sin sustento indexado"
    )
    generated_at = timezone.localtime(consultation.created_at) if consultation.created_at else timezone.now()
    signatory = getattr(consultation.user, "full_name", "") or getattr(consultation.user, "email", "") or "Dirección de Criterio Fiscal"
    signatory_role = _consultation_signatory_role(consultation)
    despacho_name = getattr(getattr(consultation.user, "despacho", None), "nombre", "") or "Despacho / corporativo del expediente"
    mark_path = _institutional_mark_path()

    page_width, page_height = letter
    content_frame = Frame(1.5 * cm, 2.05 * cm, page_width - 3 * cm, page_height - 4.0 * cm, id="legal_content")
    doc = BaseDocTemplate(
        buffer,
        pagesize=letter,
        title=f"Opinion legal preliminar CL-{consultation.id}",
        author="Materialidad Legal Fiscal",
        subject="Consulta Legal Inteligente",
    )
    doc.addPageTemplates([
        PageTemplate(id="content", frames=[content_frame], onPage=_legal_consultation_header_footer_factory(consultation_type_label)),
    ])

    content_width = page_width - 3 * cm
    elements: list[Any] = []
    elements.append(Spacer(1, 0.45 * cm))

    brand_cell: Any
    if mark_path:
        brand_cell = Table([
            [
                Image(mark_path, width=1.2 * cm, height=1.2 * cm),
                Paragraph(
                    "<b>Materialidad Legal Fiscal</b><br/><font size='8' color='#64748B'>Emision institucional</font>",
                    styles["table_cell"],
                ),
            ]
        ], colWidths=[1.6 * cm, 4.6 * cm])
        brand_cell.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
    else:
        brand_cell = Paragraph(
            "<b>Materialidad Legal Fiscal</b><br/><font size='8' color='#64748B'>Emision institucional</font>",
            styles["table_cell"],
        )

    executive_header = Table([
        [
            brand_cell,
            Paragraph(
                "<font size='17'><b>OPINION LEGAL PRELIMINAR</b></font><br/><font size='9' color='#64748B'>Sintesis ejecutiva para revision directiva. El dictamen inicia en esta misma hoja.</font>",
                styles["table_cell"],
            ),
        ]
    ], colWidths=[6.5 * cm, content_width - 6.5 * cm])
    executive_header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROUNDEDCORNERS", [8, 8, 8, 8]),
    ]))
    elements.append(executive_header)
    elements.append(Spacer(1, 10))

    summary_rows = [
        [
            Paragraph("<b>Folio</b>", styles["table_cell_bold"]),
            Paragraph(f"CL-{consultation.id}", styles["table_cell"]),
            Paragraph("<b>Fecha</b>", styles["table_cell_bold"]),
            Paragraph(generated_at.strftime("%d/%m/%Y %H:%M hrs."), styles["table_cell"]),
        ],
        [
            Paragraph("<b>Tipo</b>", styles["table_cell_bold"]),
            Paragraph(_sanitize_pdf_text(consultation_type_label), styles["table_cell"]),
            Paragraph("<b>Sustento</b>", styles["table_cell_bold"]),
            Paragraph(_sanitize_pdf_text(support_label), styles["table_cell"]),
        ],
        [
            Paragraph("<b>Clasificación</b>", styles["table_cell_bold"]),
            Paragraph("Confidencial - Uso interno", styles["table_cell"]),
            Paragraph("<b>Referencias</b>", styles["table_cell_bold"]),
            Paragraph(str(len(references_payload)), styles["table_cell"]),
        ],
    ]
    summary_table = Table(summary_rows, colWidths=[2.4 * cm, 5.2 * cm, 2.8 * cm, content_width - 10.4 * cm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), _COLOR_LIGHT_BG),
        ("BACKGROUND", (2, 0), (2, -1), _COLOR_LIGHT_BG),
        ("ROWBACKGROUNDS", (1, 0), (1, -1), [_COLOR_WHITE, colors.HexColor("#F8FAFC")]),
        ("ROWBACKGROUNDS", (3, 0), (3, -1), [_COLOR_WHITE, colors.HexColor("#F8FAFC")]),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 8))

    question_table = Table([
        [Paragraph("<b>Planteamiento sometido a revision</b>", styles["table_cell_bold"])],
        [Paragraph(_sanitize_pdf_text(consultation.question), styles["table_cell"])],
    ], colWidths=[content_width])
    question_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    elements.append(question_table)
    elements.append(Spacer(1, 8))

    if consultation.context:
        context_table = Table([
            [Paragraph("<b>Contexto operativo relevante</b>", styles["table_cell_bold"])],
            [Paragraph(_sanitize_pdf_text(consultation.context), styles["table_cell"])],
        ], colWidths=[content_width])
        context_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F8FAFC")),
            ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ]))
        elements.append(context_table)
        elements.append(Spacer(1, 8))

    guidance_table = Table([
        [
            Paragraph(
                "<b>Alcance</b><br/><font size='9'>Documento ejecutivo preliminar. Debe complementarse con revision juridica final, confirmacion de vigencia y validacion integral del expediente.</font>",
                styles["table_cell"],
            ),
            Paragraph(
                f"<b>Emitido por</b><br/><font size='9'>{_sanitize_pdf_text(signatory)}<br/>{_sanitize_pdf_text(signatory_role)}<br/>{_sanitize_pdf_text(despacho_name)}</font>",
                styles["table_cell"],
            ),
        ]
    ], colWidths=[content_width * 0.58, content_width * 0.42])
    guidance_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(guidance_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("DICTAMEN Y ANALISIS", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 6))
    elements.extend(_markdown_to_pdf_flowables(consultation.answer, styles))

    elements.append(Spacer(1, 10))
    elements.append(Paragraph("ANEXO DE REFERENCIAS", styles["section_title"]))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=_COLOR_ACCENT))
    elements.append(Spacer(1, 8))

    if not references_payload:
        elements.append(Paragraph("No se localizaron referencias indexadas suficientes para integrar un anexo documental.", styles["body"]))
    else:
        for idx, ref in enumerate(references_payload, start=1):
            title = ref.get("rubro") or ref.get("identifier") or (ref.get("articulo") and f"Artículo {ref.get('articulo')}") or ref.get("ley") or f"Referencia {idx}"
            elements.append(Paragraph(f"REFERENCIA {idx}. {_sanitize_pdf_text(str(title))}", styles["subsection_title"]))
            ref_rows = [
                [Paragraph("<b>Ordenamiento</b>", styles["table_cell_bold"]), Paragraph(_sanitize_pdf_text(ref.get("ordenamiento") or ref.get("ley")), styles["table_cell"])],
                [Paragraph("<b>Vigencia</b>", styles["table_cell_bold"]), Paragraph(_sanitize_pdf_text(ref.get("estatus_vigencia") or "Sin etiqueta de vigencia"), styles["table_cell"])],
                [Paragraph("<b>Identificación</b>", styles["table_cell_bold"]), Paragraph(_sanitize_pdf_text(" · ".join(filter(None, [ref.get("articulo") and f"Art. {ref.get('articulo')}", ref.get("fraccion") and f"Frac. {ref.get('fraccion')}", ref.get("registro_digital") and f"Registro {ref.get('registro_digital')}", ref.get("autoridad_emisora")]))), styles["table_cell"])],
            ]
            if ref.get("match_reason"):
                ref_rows.append([Paragraph("<b>Pertinencia</b>", styles["table_cell_bold"]), Paragraph(_sanitize_pdf_text(ref.get("match_reason")), styles["table_cell"])])
            ref_table = Table(ref_rows, colWidths=[4.8 * cm, content_width - 4.8 * cm])
            ref_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), _COLOR_LIGHT_BG),
                ("BOX", (0, 0), (-1, -1), 0.8, _COLOR_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, _COLOR_BORDER),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ]))
            elements.append(ref_table)
            elements.append(Spacer(1, 5))
            elements.append(Paragraph(_sanitize_pdf_text(ref.get("extracto") or ref.get("resumen")), styles["body"]))
            elements.append(Spacer(1, 10))

    elements.append(Spacer(1, 8))
    elements.append(HRFlowable(width="100%", thickness=2, color=_COLOR_PRIMARY))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Firma institucional emitida electrónicamente por Materialidad Legal Fiscal para presentación a comités y órganos de administración.", styles["legal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
